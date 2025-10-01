import os
from pathlib import Path
from typing import Dict, List, Any
from datetime import datetime
import pypdf
from docx import Document
import markdown

class DocumentProcessor:
    """
    Handles multiple document formats with metadata extraction.
    Falls back gracefully if advanced parsing fails.
    """
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt', '.md', '.html']
    
    def process_file(self, file_path: str) -> Dict[str, Any]:
        """Main entry point for document processing"""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = path.suffix.lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported format: {file_ext}")
        
        # Extract metadata
        metadata = self._extract_metadata(path)
        
        # Extract content based on file type
        if file_ext == '.pdf':
            content = self._process_pdf(path)
        elif file_ext == '.docx':
            content = self._process_docx(path)
        elif file_ext in ['.txt', '.md']:
            content = self._process_text(path)
        elif file_ext == '.html':
            content = self._process_html(path)
        else:
            content = ""
        
        return {
            "content": content,
            "metadata": metadata,
            "file_path": str(path),
            "file_name": path.name,
            "file_type": file_ext[1:],  # Remove dot
        }
    
    def _extract_metadata(self, path: Path) -> Dict[str, Any]:
        """Extract file metadata"""
        stat = path.stat()
        
        return {
            "file_name": path.name,
            "file_size": stat.st_size,
            "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
            "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
            "extension": path.suffix[1:],
        }
    
    def _process_pdf(self, path: Path) -> str:
        """Extract text from PDF with page preservation"""
        try:
            text_parts = []
            with open(path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text.strip():
                        # Add page marker for citation
                        text_parts.append(f"[PAGE {page_num}]\n{page_text}")
            
            return "\n\n".join(text_parts)
        
        except Exception as e:
            print(f"Error processing PDF {path}: {e}")
            return ""
    
    def _process_docx(self, path: Path) -> str:
        """Extract text from DOCX with structure preservation"""
        try:
            doc = Document(path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    # Detect headings
                    if para.style.name.startswith('Heading'):
                        text_parts.append(f"\n## {para.text}\n")
                    else:
                        text_parts.append(para.text)
            
            # Process tables
            for table in doc.tables:
                table_text = self._extract_table(table)
                text_parts.append(f"\n[TABLE]\n{table_text}\n")
            
            return "\n".join(text_parts)
        
        except Exception as e:
            print(f"Error processing DOCX {path}: {e}")
            return ""
    
    def _extract_table(self, table) -> str:
        """Extract table content as formatted text"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows)
    
    def _process_text(self, path: Path) -> str:
        """Process plain text and markdown files"""
        try:
            with open(path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # If markdown, preserve structure
            if path.suffix.lower() == '.md':
                # Keep markdown as is - helps with structure
                return content
            
            return content
        
        except Exception as e:
            print(f"Error processing text file {path}: {e}")
            return ""
    
    def _process_html(self, path: Path) -> str:
        """Basic HTML processing"""
        try:
            with open(path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            # Simple HTML tag removal (you can use BeautifulSoup if needed)
            import re
            text = re.sub('<[^<]+?>', '', html_content)
            return text
        
        except Exception as e:
            print(f"Error processing HTML {path}: {e}")
            return ""


# Test function
if __name__ == "__main__":
    processor = DocumentProcessor()
    
    # Test with a sample file
    result = processor.process_file("sample.pdf")
    print(f"Extracted {len(result['content'])} characters")
    print(f"Metadata: {result['metadata']}")